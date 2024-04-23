import os
import shutil
import datetime
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# Задаем исходную директорию с шаблонами
sourceDirectory = r"C:\Users\1C\Desktop\Шаблоны"

# Пути к различным шаблонам
template_paths = [
    os.path.join(sourceDirectory, "SPA салон"),
    os.path.join(sourceDirectory, "Салон"),
    os.path.join(sourceDirectory, "Стоматология")
]

def validate_fields():
    required_entries = [
        target_folder_path,
        old_version_1, old_version_2, old_version_3, old_version_4,
        new_version_1, new_version_2, new_version_3, new_version_4,
        platformversion_1, platformversion_2, platformversion_3, platformversion_4,
        slk_1, slk_2, slk_3, slk_4
    ]
    # Проверяем, что все обязательные поля заполнены
    for entry in required_entries:
        if not entry.get():
            return False

    return True

def replace_text_in_file(file_path, replacements):
    """ Замена текста в файле """
    with open(file_path, 'r', encoding='latin1') as file:
        content = file.read()
    
    for search_text, replace_text in replacements.items():
        content = content.replace(search_text, replace_text)
    
    with open(file_path, 'w', encoding='latin1') as file:
        file.write(content)

def replace_text_in_docx_file(file_path, replacements):
    """ Замена текста в документе .docx """
    doc = Document(file_path)

    for search_text, replace_text in replacements.items():
        # Замена текста в параграфах документа
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                paragraph.text = paragraph.text.replace(search_text, replace_text)

        # Замена текста в ячейках таблиц документа
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if search_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(search_text, replace_text)

    doc.save(file_path)

def copy_files_with_replacement(source_dir, target_dir, replacements):
    """ Копирование файлов с заменой текста """
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            source_file_path = os.path.join(root, file)
            relative_path = os.path.relpath(source_file_path, source_dir)
            target_file_path = os.path.join(target_dir, relative_path)

            os.makedirs(os.path.dirname(target_file_path), exist_ok=True)

            if file.endswith('.txt') or file.endswith('.docx'):
                shutil.copyfile(source_file_path, target_file_path)

                if file.endswith('.txt'):
                    replace_text_in_file(target_file_path, replacements)
                elif file.endswith('.docx'):
                    replace_text_in_docx_file(target_file_path, replacements)

def browse_folder():
    """ Выбор целевой папки """
    folder_path = filedialog.askdirectory(title="Выберите целевую папку")
    if folder_path:
        target_folder_path.delete(0, tk.END)
        target_folder_path.insert(0, folder_path)

def delete_specific_files_in_directory(directory_path, files_to_delete):
    """ Рекурсивно удаляет указанные файлы из целевой директории и всех её подпапок """
    for root, dirs, files in os.walk(directory_path):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            if file_name in files_to_delete:
                os.remove(file_path)

def update_files():
    """ Обновление файлов """
    if not validate_fields():
        print("Ошибка", "Заполните все поля перед обновлением файлов")
        return
    
    try:
        target_path = target_folder_path.get()
        selected_template_index = template_combobox.current()
        source_directory = template_paths[selected_template_index]

        # Формирование старой и новой версий
        old_version1 = f"{old_version_1.get()}.{old_version_2.get()}.{old_version_3.get()}"
        old_version2 = f"{old_version_1.get()}_{old_version_2.get()}_{old_version_3.get()}_{old_version_4.get()}"
        old_version3 = f"{old_version_1.get()}.{old_version_2.get()}.{old_version_3.get()}.{old_version_4.get()}"

        new_version1 = f"{new_version_1.get()}.{new_version_2.get()}.{new_version_3.get()}"
        new_version2 = f"{new_version_1.get()}_{new_version_2.get()}_{new_version_3.get()}_{new_version_4.get()}"
        new_version3 = f"{new_version_1.get()}.{new_version_2.get()}.{new_version_3.get()}.{new_version_4.get()}"

        platform_version_1 = f"{platformversion_1.get()}.{platformversion_2.get()}.{platformversion_3.get()}"
        platform_version_2 = f"{platformversion_1.get()}.{platformversion_2.get()}.{platformversion_3.get()}.{platformversion_4.get()}"

        slk_version = f"{slk_1.get()}.{slk_2.get()}.{slk_3.get()}.{slk_4.get()}"

        current_year = datetime.datetime.now().year
        current_date = datetime.datetime.now().strftime('%d.%m.%Y')

        replacements = {
            'oldversion_1': old_version1,
            'oldversion_2': old_version2,
            'oldversion_3': old_version3,
            'newversion_1': new_version1,
            'newversion_2': new_version2,
            'newversion_3': new_version3,
            'platformversion_1': platform_version_1,
            'platformversion_2': platform_version_2,
            'slkversion': slk_version,
            'date_year': str(current_year),
            'date_today': current_date
        }

        # Копирование файлов с заменой текста
        copy_files_with_replacement(source_directory, target_path, replacements)

        # Удаляем определенные файлы из целевой папки
        files_to_delete = ['1cv8.efd', 'setup.exe', 'setup', '1Cv8.dt']
        delete_specific_files_in_directory(target_path, files_to_delete)

        print("Обновление файлов", "Файлы успешно обновлены")
    except Exception as e:
        print("Ошибка", f"Произошла ошибка: {e}")

# Создание основного окна
root = tk.Tk()
root.title("Обновление файлов")

# Фрейм для целевой папки
frame_target = tk.Frame(root)
frame_target.pack(pady=10)

label_target = tk.Label(frame_target, text="Путь к целевой папке:")
label_target.grid(row=0, column=0, padx=10, sticky="w")

target_folder_path = tk.Entry(frame_target, width=50)
target_folder_path.grid(row=0, column=1, padx=10, pady=5, sticky="we")

browse_button = tk.Button(frame_target, text="Обзор", command=browse_folder)
browse_button.grid(row=0, column=2, padx=10, pady=5)

# Фрейм для выбора шаблона
frame_template = tk.Frame(root)
frame_template.pack(pady=10)

label_template = tk.Label(frame_template, text="Выберите шаблон:")
label_template.grid(row=0, column=0, padx=10, sticky="w")

template_combobox = ttk.Combobox(frame_template, values=["SPA салон", "Салон", "Стоматология"])
template_combobox.grid(row=0, column=1, padx=10, pady=5)
template_combobox.current(0)

# Фрейм для старой версии
frame_old_version = tk.Frame(root)
frame_old_version.pack(pady=10)

label_old_version = tk.Label(frame_old_version, text="Старая версия (Пример: 3 0 32 1):", fg="red")
label_old_version.grid(row=0, column=0, padx=10, sticky="w")

old_version_1 = tk.Entry(frame_old_version, width=3)
old_version_1.grid(row=0, column=1, padx=5, pady=5, sticky="w")
old_version_2 = tk.Entry(frame_old_version, width=3)
old_version_2.grid(row=0, column=2, padx=5, pady=5, sticky="w")
old_version_3 = tk.Entry(frame_old_version, width=3)
old_version_3.grid(row=0, column=3, padx=5, pady=5, sticky="w")
old_version_4 = tk.Entry(frame_old_version, width=3)
old_version_4.grid(row=0, column=4, padx=5, pady=5, sticky="w")

# Фрейм для новой версии
frame_new_version = tk.Frame(root)
frame_new_version.pack(pady=10)

label_new_version = tk.Label(frame_new_version, text="Новая версия (Пример: 3 0 33 1):", fg="blue")
label_new_version.grid(row=0, column=0, padx=10, sticky="w")

new_version_1 = tk.Entry(frame_new_version, width=3)
new_version_1.grid(row=0, column=1, padx=5, pady=5, sticky="w")
new_version_2 = tk.Entry(frame_new_version, width=3)
new_version_2.grid(row=0, column=2, padx=5, pady=5, sticky="w")
new_version_3 = tk.Entry(frame_new_version, width=3)
new_version_3.grid(row=0, column=3, padx=5, pady=5, sticky="w")
new_version_4 = tk.Entry(frame_new_version, width=3)
new_version_4.grid(row=0, column=4, padx=5, pady=5, sticky="w")

# Фрейм для версии платформы
frame_platformversion = tk.Frame(root)
frame_platformversion.pack(pady=10)

label_platformversion = tk.Label(frame_platformversion, text="Версия платформы:", fg="green")
label_platformversion.grid(row=0, column=0, padx=10, sticky="w")

platformversion_1 = tk.Entry(frame_platformversion, width=3)
platformversion_1.grid(row=0, column=1, padx=5, pady=5, sticky="w")
platformversion_2 = tk.Entry(frame_platformversion, width=3)
platformversion_2.grid(row=0, column=2, padx=5, pady=5, sticky="w")
platformversion_3 = tk.Entry(frame_platformversion, width=3)
platformversion_3.grid(row=0, column=3, padx=5, pady=5, sticky="w")
platformversion_4 = tk.Entry(frame_platformversion, width=6)
platformversion_4.grid(row=0, column=4, padx=5, pady=5, sticky="w")

# Значения по умолчанию для версии платформы
default_platform_version = ('8', '3', '21', '1895')
platformversion_1.insert(0, default_platform_version[0])
platformversion_2.insert(0, default_platform_version[1])
platformversion_3.insert(0, default_platform_version[2])
platformversion_4.insert(0, default_platform_version[3])

# Фрейм для версии СЛК
frame_slk = tk.Frame(root)
frame_slk.pack(pady=10)

label_slk = tk.Label(frame_slk, text="Версия СЛК:", fg="green")
label_slk.grid(row=0, column=0, padx=10, sticky="w")

slk_1 = tk.Entry(frame_slk, width=3)
slk_1.grid(row=0, column=1, padx=5, pady=5, sticky="w")
slk_2 = tk.Entry(frame_slk, width=3)
slk_2.grid(row=0, column=2, padx=5, pady=5, sticky="w")
slk_3 = tk.Entry(frame_slk, width=3)
slk_3.grid(row=0, column=3, padx=5, pady=5, sticky="w")
slk_4 = tk.Entry(frame_slk, width=6)
slk_4.grid(row=0, column=4, padx=5, pady=5, sticky="w")

# Значения по умолчанию для версии СЛК
default_slk_version = ('3', '0', '33', '11307')
slk_1.insert(0, default_slk_version[0])
slk_2.insert(0, default_slk_version[1])
slk_3.insert(0, default_slk_version[2])
slk_4.insert(0, default_slk_version[3])

# Update button
update_button = tk.Button(root, text="Обновить файлы", command=update_files)
update_button.pack(pady=20)

def bind_validation(event, entry):
    validate_fields()

for entry in [
    target_folder_path, old_version_1, old_version_2, old_version_3, old_version_4,
    new_version_1, new_version_2, new_version_3, new_version_4,
    platformversion_1, platformversion_2, platformversion_3, platformversion_4,
    slk_1, slk_2, slk_3, slk_4
]:
    entry.bind("<FocusOut>", lambda event, entry=entry: bind_validation(event, entry))


root.mainloop()